from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs"
OUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "Nasaq_SRS_AR_v1.0.docx"
LOGO = ROOT / "public" / "nasaq-logo.png"

GREEN = "0B5D45"
GREEN_DARK = "083B2D"
GREEN_LIGHT = "E8F3EF"
GOLD = "B88A22"
INK = "17243A"
MUTED = "5B6678"
BORDER = "D7E0DD"
LIGHT = "F5F7F6"
WHITE = "FFFFFF"
RED = "9E2A2B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=10.5, color=INK, bold=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_para(doc, text="", *, size=10.5, color=INK, bold=False, after=6, before=0,
             align=WD_ALIGN_PARAGRAPH.RIGHT, keep=False):
    p = doc.add_paragraph()
    set_rtl(p, align)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = keep
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_rtl(p)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, color=GREEN, bold=True)
    elif level == 2:
        set_font(run, size=13, color=GREEN_DARK, bold=True)
    else:
        set_font(run, size=11.5, color=INK, bold=True)
    return p


def add_list(doc, items, *, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for text in items:
        p = doc.add_paragraph(style=style)
        set_rtl(p)
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_font(run)


def add_callout(doc, title, body, fill=GREEN_LIGHT, title_color=GREEN_DARK):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    right = OxmlElement("w:right")
    right.set(qn("w:val"), "single")
    right.set(qn("w:sz"), "18")
    right.set(qn("w:space"), "6")
    right.set(qn("w:color"), title_color)
    p_bdr.append(right)
    p_pr.append(p_bdr)
    r = p.add_run(title)
    set_font(r, size=11, color=title_color, bold=True)
    r.add_break()
    r2 = p.add_run(body)
    set_font(r2, size=10.2)


def add_table(doc, headers, rows, widths, *, header_fill=GREEN_DARK, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9.2, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_index % 2:
                set_cell_shading(cells[idx], LIGHT)
            p = cells[idx].paragraphs[0]
            set_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT if idx else WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value))
            set_font(r, size=font_size, color=INK, bold=(idx == 0))
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_para(doc, "", after=2)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("صفحة ")
    set_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, GREEN, 16, 8),
        ("Heading 2", 13, GREEN_DARK, 12, 6),
        ("Heading 3", 11.5, INK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(10.5)

    header = section.header
    hp = header.paragraphs[0]
    set_rtl(hp)
    hp.paragraph_format.space_after = Pt(3)
    run = hp.add_run("نَسَق | مواصفات متطلبات البرمجيات")
    set_font(run, size=8.5, color=MUTED, bold=True)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_page_number(section.footer.paragraphs[0])


def cover(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        inline = run.add_picture(str(LOGO), width=Inches(1.05))
        try:
            inline._inline.docPr.set("descr", "شعار منصة نَسَق لإدارة المخيمات")
        except Exception:
            pass
    add_para(doc, "نَسَق لإدارة المخيمات والاستجابة الإنسانية", size=13, color=GREEN, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, "مواصفات متطلبات البرمجيات", size=25, color=INK, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, "Software Requirements Specification (SRS)", size=12, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_callout(doc, "حالة الوثيقة", "نسخة الإصدار 1.0 - مرجع الإطلاق والتشغيل وقبول المنتج النهائي.")
    add_table(doc, ["القيمة", "البيان"], [
        ("1.0", "الإصدار"),
        ("20 أغسطس 2026", "تاريخ الإصدار"),
        ("جاهز للمراجعة والاعتماد", "الحالة"),
        ("مالك النظام / الجهات المشغلة / فرق التقنية والدعم", "الجمهور"),
        ("Next.js 16 + Supabase PostgreSQL", "المنصة التقنية"),
    ], [6500, 2860], header_fill=GREEN_DARK, font_size=9.5)
    add_para(doc, "تصنيف الوثيقة: تشغيلي وتقني - للاستخدام المؤسسي", size=9, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=16)
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_document(doc)
    cover(doc)

    add_heading(doc, "1. ضبط الوثيقة", 1)
    add_table(doc, ["التفصيل", "البند"], [
        ("تحديد السلوك المطلوب للنظام ومعايير قبوله قبل النشر المؤسسي.", "الغرض"),
        ("لوحات مدير المخيم والمشرف العام، البيانات، المصادقة، التصدير والطباعة والتشغيل.", "النطاق"),
        ("الكود الحالي، مخطط Supabase، اختبارات العقود والبناء الإنتاجي وفحص الواجهات.", "مصادر المواصفة"),
        ("يعتمد عند اكتمال عناصر الإطلاق الخارجية المذكورة في القسم 12.", "قاعدة الاعتماد"),
    ], [6700, 2660])
    add_heading(doc, "1.1 سجل الإصدارات", 2)
    add_table(doc, ["موجز التغيير", "التاريخ", "الإصدار"], [
        ("إصدار أساسي بعد تقوية الأمان وعزل بيانات المخيمات وتحسين الهاتف.", "20 أغسطس 2026", "1.0"),
    ], [5600, 2300, 1460])
    add_heading(doc, "1.2 مصطلحات أساسية", 2)
    add_table(doc, ["التعريف", "المصطلح"], [
        ("معرّف المخيم المنطقي المستخدم في الربط والعزل بين السجلات.", "Camp ID"),
        ("المستخدم الأعلى صلاحية والمسؤول عن إنشاء المخيمات والاشتراكات والحذف النهائي.", "المشرف العام"),
        ("مستخدم مؤسسي محصور في بيانات مخيم واحد.", "مدير المخيم"),
        ("سياسات PostgreSQL التي تمنع الوصول إلى صفوف مخيم آخر.", "RLS"),
        ("حذف المخيم وجميع السجلات التابعة له ضمن عملية قاعدة بيانات ذرّية.", "الحذف المتسلسل"),
    ], [6800, 2560])

    add_heading(doc, "2. الملخص التنفيذي", 1)
    add_para(doc, "نَسَق منصة ويب متعددة المخيمات لتنظيم بيانات الأسر والترشيحات وإدارة حسابات المخيمات والاشتراكات. تفصل المنصة كل مخيم عن غيره في طبقة قاعدة البيانات، وتوفر للمشرف العام إدارة مركزية دون تخزين كلمات مرور داخل جداول الأعمال.")
    add_callout(doc, "قرار هندسي أساسي", "Supabase هو مصدر الحقيقة الوحيد في وضع الإنتاج. لا يعرض النظام بيانات تجريبية على أنها بيانات حقيقية عند تعذر الاتصال، بل يصرّح بالخطأ للمستخدم ويمنع نجاح العمليات الوهمي.")
    add_heading(doc, "2.1 الأهداف", 2)
    add_list(doc, [
        "توحيد سجل الأسر والترشيحات لكل مخيم ضمن واجهة عربية واضحة ومتجاوبة.",
        "حماية بيانات كل مخيم بسياسات عزل مطبقة داخل PostgreSQL وليس في الواجهة فقط.",
        "توفير كشوفات رسمية قابلة للطباعة وقوالب Excel فارغة للاستيراد المنظم.",
        "إتاحة إدارة مركزية للاشتراكات والإعلانات وحسابات مديري المخيمات.",
        "تمكين النشر على استضافة تدعم Next.js مع أسرار خادم منفصلة عن مفاتيح المتصفح.",
    ])
    add_heading(doc, "2.2 خارج النطاق الحالي", 2)
    add_list(doc, [
        "تطبيق جوال أصلي منفصل؛ الواجهة الحالية تطبيق ويب متجاوب.",
        "العمل دون اتصال ثم المزامنة اللاحقة؛ لا يوجد نجاح محلي بديل في الإنتاج.",
        "سجل تدقيق غير قابل للتعديل لكل تغيير؛ يعد تحسينًا مستقبليًا موصى به.",
        "إرسال رسائل SMS أو بريد إلكتروني تلقائي؛ غير مضمّن في الإصدار الحالي.",
    ])
    doc.add_page_break()

    add_heading(doc, "3. أصحاب المصلحة والأدوار", 1)
    add_table(doc, ["المسؤوليات", "الصلاحية", "الدور"], [
        ("إنشاء وتعديل وحذف المخيمات، إدارة الاشتراك، نشر الإعلانات، مراجعة طلبات التجديد.", "كل المخيمات", "المشرف العام"),
        ("إدارة ملف المخيم والأسر والترشيحات والتقارير ضمن المخيم المرتبط بالحساب.", "مخيم واحد", "مدير المخيم"),
        ("النسخ الاحتياطي، المتغيرات السرية، المراقبة والاستجابة للحوادث.", "بيئة التشغيل", "مسؤول التقنية"),
        ("اعتماد الغرض من جمع البيانات وسياسة الاحتفاظ والاستجابة لطلبات أصحاب البيانات.", "حوكمة البيانات", "الجهة المالكة"),
    ], [5100, 1900, 2360], font_size=8.8)
    add_heading(doc, "3.1 مصفوفة الصلاحيات", 2)
    add_table(doc, ["مدير المخيم", "المشرف العام", "الوظيفة"], [
        ("قراءة/تعديل مخيمه", "قراءة/تعديل الكل", "ملف المخيم"),
        ("ضمن مخيمه", "إشراف عند الحاجة", "الأسر والترشيحات"),
        ("لا", "نعم", "إنشاء مخيم وحساب مدير"),
        ("لا", "نعم مع تأكيد المعرّف", "الحذف النهائي"),
        ("طلب تجديد", "موافقة/رفض", "الاشتراك"),
        ("عرض", "إنشاء وإدارة", "الإعلانات"),
    ], [2500, 2500, 4360])
    add_callout(doc, "قاعدة منع التصعيد", "لا يجوز الاعتماد على user_metadata لاتخاذ قرار صلاحية. يشتق الدور والارتباط بالمخيم من بيانات خادم موثوقة وسياسات RLS.")

    add_heading(doc, "4. نظرة عامة على النظام", 1)
    add_table(doc, ["المسؤولية", "الطبقة"], [
        ("Next.js App Router، صفحات عربية RTL، مكونات متجاوبة، استيراد وتصدير.", "واجهة الويب"),
        ("مسارات API إدارية محمية لإنشاء وتعديل وحذف الحسابات والمخيمات.", "طبقة الخادم"),
        ("Supabase Auth للجلسات؛ PostgreSQL للبيانات؛ RLS للعزل؛ Realtime للتحديثات المحددة.", "الخدمات المُدارة"),
        ("Excel للاستيراد والتصدير؛ طباعة المتصفح لإنتاج PDF رسمي.", "التكاملات"),
    ], [6700, 2660])
    add_heading(doc, "4.1 حدود الثقة", 2)
    add_list(doc, [
        "المفتاح القابل للنشر فقط مسموح في المتصفح.",
        "المفتاح السري يبقى على الخادم ويستخدم في المسارات الإدارية المحمية فقط.",
        "كل عملية بيانات تتحقق من خطأ Supabase ومن الصفوف المتأثرة قبل إعلان النجاح.",
        "RLS ومفاتيح الربط الخارجية تبقيان نافذتين حتى لو تم العبث بطلبات المتصفح.",
    ])
    doc.add_page_break()

    add_heading(doc, "5. المتطلبات الوظيفية", 1)
    add_heading(doc, "5.1 المصادقة والجلسات", 2)
    auth_rows = [
        ("FR-001", "يسجل المستخدم الدخول باسم المستخدم وكلمة المرور من صفحة واحدة.", "جلسة صالحة وتحويل إلى اللوحة المناسبة."),
        ("FR-002", "يحدد الخادم دور المستخدم ومخيمه بعد نجاح Supabase Auth.", "لا يعتمد القرار على بيانات وصفية قابلة لتعديل المستخدم."),
        ("FR-003", "يمنع غير المسجل من المسارات المحمية.", "إعادة توجيه إلى /login أو رفض API."),
        ("FR-004", "يسجل الخروج ويمحو جلسة التطبيق.", "لا تبقى صفحة محمية متاحة بعد الخروج."),
        ("FR-005", "يعرض فشل الدخول برسالة واضحة دون كشف سبب أمني تفصيلي.", "لا يكشف وجود المستخدم أو كلمة المرور."),
    ]
    add_table(doc, ["معيار القبول", "المتطلب", "المعرف"], auth_rows, [3500, 4400, 1460], font_size=8.5)

    add_heading(doc, "5.2 إدارة المخيمات والحسابات", 2)
    camp_rows = [
        ("FR-006", "ينشئ المشرف العام مخيمًا وحساب مدير ضمن عملية منسقة.", "إنشاء Auth ثم RPC ذرّي للملف والربط."),
        ("FR-007", "يتحقق النظام من معرّف واسم المستخدم ومن سياسة كلمة المرور.", "10 أحرف على الأقل مع حرف ورقم."),
        ("FR-008", "يعدل المشرف بيانات المخيم والاشتراك وحساب المدير.", "تنعكس التغييرات بعد إعادة القراءة من المصدر."),
        ("FR-009", "يحذف المشرف المخيم نهائيًا بعد كتابة Camp ID للتأكيد.", "حذف السجلات التابعة وقفل الوصول فورًا."),
        ("FR-010", "يعرض المشرف عدد المخيمات وحالة الاشتراكات والوقت المتبقي.", "الحالة مشتقة من تاريخ الانتهاء."),
        ("FR-011", "يبحث المشرف بالاسم أو المدير أو المعرّف.", "تتحدث النتائج دون إعادة تحميل الصفحة."),
    ]
    add_table(doc, ["معيار القبول", "المتطلب", "المعرف"], camp_rows, [3500, 4400, 1460], font_size=8.5)
    doc.add_page_break()

    add_heading(doc, "5.3 سجل الأسر", 2)
    family_rows = [
        ("FR-012", "ينشئ مدير المخيم سجل أسرة مرتبطًا بمخيمه.", "لا يقبل السجل دون camp_id الموثوق."),
        ("FR-013", "يعدل ويحذف السجل مع إظهار خطأ قاعدة البيانات عند الفشل.", "لا يظهر نجاح وهمي أو بديل محلي."),
        ("FR-014", "يعرض ويبحث ويرشح السجلات.", "النتيجة محصورة بالمخيم الحالي."),
        ("FR-015", "يستورد ملف Excel مع معاينة والتحقق قبل الاعتماد.", "رفض الصفوف غير الصالحة وبيان سببها."),
        ("FR-016", "يصدر كشف Excel ويجهز نسخة طباعة باسم وشعار ومسؤول المخيم.", "ترويسة رسمية وحقول واضحة للطباعة."),
        ("FR-017", "يوفر قالب أسر فارغًا.", "ملف XLSX ثابت قابل للتعبئة ثم الرفع."),
    ]
    add_table(doc, ["معيار القبول", "المتطلب", "المعرف"], family_rows, [3500, 4400, 1460], font_size=8.5)
    add_para(doc, "الحقول الأساسية: اسم رب الأسرة، رقم الهوية، تاريخ الميلاد، الحالة الاجتماعية، بيانات الزوجة عند الحاجة، الهاتف، عدد الأفراد، الموقع، والملاحظات.", color=MUTED)

    add_heading(doc, "5.4 سجل الترشيحات", 2)
    nomination_rows = [
        ("FR-018", "ينشئ ويعدل ويحذف مدير المخيم الترشيحات التفصيلية.", "تخزن الحقول التفصيلية ولا تختزل في حقول عامة."),
        ("FR-019", "يدعم بيانات الزوجة الأولى والثانية عند انطباقها.", "حقول الاسم والهوية مستقلة وقابلة للتحقق."),
        ("FR-020", "يدعم توزيع أفراد الأسرة حسب العمر والجنس.", "0-2، 3-5، 6-18، 19-60، وأكثر من 60."),
        ("FR-021", "يسجل مؤشرات الهشاشة.", "إعاقة، مرض مزمن، حمل/رضاعة، أسرة تعيلها امرأة."),
        ("FR-022", "يسجل العنوان الحالي والأصلي والمحافظة وبيانات مركز الإيواء.", "تخزن بيانات الموقع والاتصال والإحداثيات."),
        ("FR-023", "يستورد ويصدر الترشيحات ويوفر قالبًا فارغًا.", "المعاينة والتحقق يسبقان الحفظ."),
    ]
    add_table(doc, ["معيار القبول", "المتطلب", "المعرف"], nomination_rows, [3500, 4400, 1460], font_size=8.5)
    doc.add_page_break()

    add_heading(doc, "5.5 اللوحة والتقارير والإعدادات", 2)
    dashboard_rows = [
        ("FR-024", "تعرض لوحة المخيم مؤشرات الأسر والأفراد والترشيحات والحالات الهشة.", "الأرقام محسوبة من بيانات المخيم الحالي."),
        ("FR-025", "يعرض النظام إعلانًا عامًا نشطًا.", "يختفي عند التعطيل أو انتهاء السياق."),
        ("FR-026", "يعدل مدير المخيم الاسم والمنطقة والمدير والهاتف والشعار.", "الحفظ في Supabase وإعادة القراءة."),
        ("FR-027", "يطلب مدير المخيم تجديد الاشتراك.", "يظهر الطلب للمشرف العام بحالة واضحة."),
        ("FR-028", "يعتمد المشرف طلب التجديد عبر RPC مقيدة الصلاحية.", "تحديث الطلب والاشتراك باتساق."),
        ("FR-029", "تنقل صفحة الطباعة بيانات الكشف عبر sessionStorage ثم تمسحها.", "لا تبقى بيانات الكشف الدورية في localStorage."),
        ("FR-030", "تعمل الوظائف الأساسية باللمس ولوحة المفاتيح.", "هدف لمس لا يقل عن 44px وتركيز مرئي."),
    ]
    add_table(doc, ["معيار القبول", "المتطلب", "المعرف"], dashboard_rows, [3500, 4400, 1460], font_size=8.5)

    add_heading(doc, "6. نموذج البيانات", 1)
    add_table(doc, ["العلاقة والغرض", "الكيان"], [
        ("السجل الجذري للمؤسسة؛ الحذف يتسلسل إلى بيانات المخيم.", "camps"),
        ("يربط auth.users بالدور والمخيم؛ لا يحتوي كلمة مرور.", "users"),
        ("سجل الأسر، مرتبط بالمخيم بمفتاح خارجي.", "families"),
        ("سجل الترشيحات التفصيلي، مرتبط بالمخيم.", "nominations"),
        ("طلبات تجديد الاشتراك وحالاتها.", "renewal_requests"),
        ("إعلانات عامة نشطة/غير نشطة.", "announcements"),
        ("إعدادات النظام المركزية المحمية.", "system_settings"),
    ], [6700, 2660])
    add_heading(doc, "6.1 قواعد التكامل", 2)
    add_list(doc, [
        "كل سجل أسرة أو ترشيح يحمل camp_id صالحًا يشير إلى camps.",
        "حذف مخيم يزيل تلقائيًا ملفات المستخدمين والأسر والترشيحات وطلبات التجديد التابعة.",
        "لا يوجد عمود لكلمة المرور أو تجزئتها داخل public schema.",
        "المعرفات الجديدة تستخدم UUID/معرفات عشوائية مقاومة للتصادم، لا الوقت وحده.",
        "Realtime محصور في families وnominations وannouncements لتقليل سطح البيانات.",
    ])
    doc.add_page_break()

    add_heading(doc, "7. تدفقات العمل الحرجة", 1)
    add_heading(doc, "7.1 إنشاء مخيم", 2)
    add_list(doc, [
        "يدخل المشرف العام بيانات المخيم وحساب المدير وسياسة الاشتراك.",
        "يتحقق الخادم من الجلسة والدور ومن صحة الحقول وقوة كلمة المرور.",
        "ينشئ الخادم مستخدم Supabase Auth.",
        "ينفذ RPC لإنشاء ملف المخيم وربط المدير بصورة ذرّية.",
        "عند فشل قاعدة البيانات يُحذف مستخدم Auth التعويضي ولا يُعلن النجاح.",
        "تعاد قراءة قائمة المخيمات من Supabase وتظهر النتيجة للمستخدم.",
    ], numbered=True)
    add_heading(doc, "7.2 حذف مخيم نهائيًا", 2)
    add_list(doc, [
        "يفتح المشرف نافذة حذف صريحة ويكتب Camp ID كاملًا.",
        "يتحقق الخادم من الدور ويحدد حسابات Auth التابعة.",
        "ينفذ RPC الحذف المتسلسل في قاعدة البيانات أولًا.",
        "يزول وصول مدير المخيم فور حذف ملف الصلاحية وسياسات RLS.",
        "يحاول الخادم تنظيف حسابات Auth ويعيد تحذيرًا إن تعذر جزء غير حرج.",
        "تعاد القائمة من المصدر للتأكد من عدم رجوع المخيم بعد تحديث الصفحة.",
    ], numbered=True)
    add_heading(doc, "7.3 استيراد كشف", 2)
    add_list(doc, [
        "ينزل المستخدم القالب الرسمي أو يختار ملفًا متوافقًا.",
        "يحلل المتصفح الملف محليًا ويعرض معاينة الصفوف.",
        "تتحقق الحقول الإلزامية والصيغ والتكرارات.",
        "يعتمد المستخدم الاستيراد؛ تحفظ الصفوف الصالحة مع camp_id الخاص بالجلسة.",
        "يعرض النظام النتيجة الفعلية والأخطاء القابلة للتصحيح.",
    ], numbered=True)

    add_heading(doc, "8. الواجهات والتكاملات", 1)
    add_table(doc, ["الوصف", "المسار/الخدمة"], [
        ("دخول موحد وتوجيه بحسب الدور.", "/login"),
        ("لوحة مدير المخيم والمؤشرات.", "/"),
        ("إدارة الأسر والترشيحات والإعدادات.", "/families، /nominations، /settings"),
        ("صفحة طباعة مهيأة للكشوف الرسمية.", "/print"),
        ("إدارة مركزية للمخيمات والاشتراكات.", "/super-admin"),
        ("إنشاء/تعديل/حذف مخيم وحساب مدير؛ تتطلب مشرفًا عامًا.", "/api/admin/*"),
        ("Auth، Database، RLS، Realtime.", "Supabase"),
    ], [6400, 2960], font_size=9)
    doc.add_page_break()

    add_heading(doc, "9. المتطلبات غير الوظيفية", 1)
    nfr_rows = [
        ("NFR-001", "الأمان", "تطبيق RLS على كل جدول مكشوف وتقييد RPC ومنع الأسرار في العميل."),
        ("NFR-002", "الخصوصية", "جمع الحد الأدنى، سياسة احتفاظ، حق التصحيح والحذف، أجهزة موثوقة."),
        ("NFR-003", "التوافق", "دعم عروض 320 و375 و414 و768 بكسل دون تمرير أفقي."),
        ("NFR-004", "الإتاحة", "تنقل بلوحة المفاتيح، تركيز مرئي، تسميات حقول، رسائل حالة واضحة."),
        ("NFR-005", "الأداء", "تقسيم حزم الاستيراد الثقيلة وتحميلها عند الطلب فقط."),
        ("NFR-006", "سلامة البيانات", "رفض نجاح العمليات عند خطأ Supabase والتحقق من الصفوف المتأثرة."),
        ("NFR-007", "القابلية للصيانة", "مخطط SQL ومهاجرات واختبارات عقود قابلة لإعادة التشغيل."),
        ("NFR-008", "التوفر", "مراقبة أخطاء الخادم والعميل وخطة نسخ احتياطي واستعادة مجربة."),
        ("NFR-009", "الرصد", "سجلات دون كلمات مرور أو مفاتيح أو بيانات شخصية زائدة."),
        ("NFR-010", "التعريب", "واجهة RTL عربية واتساق المصطلحات وتواريخ مفهومة للمستخدم المحلي."),
    ]
    add_table(doc, ["المطلب", "المجال", "المعرف"], [(c, b, a) for a, b, c in nfr_rows], [6000, 1900, 1460], font_size=8.8)

    add_heading(doc, "9.1 ترويسات الحماية", 2)
    add_list(doc, [
        "Content-Security-Policy للحد من مصادر السكربتات والاتصالات.",
        "Strict-Transport-Security على الإنتاج عبر HTTPS.",
        "X-Content-Type-Options: nosniff وX-Frame-Options: DENY.",
        "Referrer-Policy وPermissions-Policy وCross-Origin-Opener-Policy.",
    ])
    add_callout(doc, "ملاحظة أمان قبل الإطلاق", "يجب تدوير أي مفتاح سري سبق مشاركته في محادثة أو شاشة، وتفعيل حماية كلمات المرور المسرّبة من لوحة Supabase Auth.", fill="FFF4E1", title_color=GOLD)

    add_heading(doc, "10. بيئات التشغيل والنشر", 1)
    add_table(doc, ["الاستخدام", "المتغير"], [
        ("عنوان مشروع Supabase.", "NEXT_PUBLIC_SUPABASE_URL"),
        ("المفتاح القابل للنشر في المتصفح.", "NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY"),
        ("مفتاح خادم سري لمسارات الإدارة فقط.", "SUPABASE_SECRET_KEY"),
        ("سر جلسات خادم قوي وعشوائي.", "JWT_SECRET"),
        ("رابط HTTPS النهائي للنظام.", "NEXT_PUBLIC_APP_URL"),
    ], [6100, 3260], font_size=9)
    add_heading(doc, "10.1 معايير بيئة الإنتاج", 2)
    add_list(doc, [
        "Node.js 20.9 أو أحدث، وبناء npm run build ناجح على Next.js 16.",
        "نطاق رسمي وشهادة HTTPS وإعادة توجيه HTTP إلى HTTPS.",
        "أسرار منفصلة لكل بيئة ولا توجد ملفات .env في المستودع.",
        "نسخ احتياطية لقاعدة البيانات واختبار استعادة دوري.",
        "مراقبة للأخطاء والتوافر وخطة استجابة للحوادث ودعم المستخدمين.",
    ])
    doc.add_page_break()

    add_heading(doc, "11. التحقق ومعايير القبول", 1)
    add_table(doc, ["النتيجة المطلوبة", "الفحص", "المعرف"], [
        ("نجاح جميع اختبارات العقد.", "اختبارات Node للعقود الأمنية والبيانية والمتجاوبة.", "AT-001"),
        ("خروج ناجح دون أخطاء.", "بناء Next.js الإنتاجي.", "AT-002"),
        ("لا تمرير أفقي وأهداف لمس 44px فأكثر.", "فحص 320/375/414/768.", "AT-003"),
        ("307 إلى /login.", "طلب مسار محمي دون جلسة.", "AT-004"),
        ("رفض الطلب.", "استدعاء API إداري دون مشرف عام.", "AT-005"),
        ("اختفاء المخيم وكل السجلات التابعة وعدم عودته.", "اختبار تكامل الحذف المتسلسل.", "AT-006"),
        ("صفر أعمدة كلمة مرور عامة.", "فحص مخطط public.", "AT-007"),
        ("لا تحذيرات RLS/وظائف عالية الخطورة.", "Supabase Security Advisor.", "AT-008"),
        ("تسجيل دخول وCRUD وتصدير وطباعة بنجاح.", "اختبار رحلة مستخدم على بيئة نشر متصلة.", "AT-009"),
        ("استعادة ناجحة ضمن زمن متفق عليه.", "اختبار نسخة احتياطية واستعادة.", "AT-010"),
    ], [3200, 4700, 1460], font_size=8.5)
    add_heading(doc, "11.1 قرار القبول", 2)
    add_para(doc, "يعد الإصدار مقبولًا تقنيًا للنشر عندما تنجح AT-001 إلى AT-008، وتعد الجهة جاهزة لاستقبال بيانات حقيقية بعد تنفيذ AT-009 وAT-010 على بيئة الإنتاج وتوثيق النتيجة.")

    add_heading(doc, "12. حالة الإصدار ومخاطر الإطلاق", 1)
    add_table(doc, ["الإجراء", "الحالة", "البند"], [
        ("مطبق ومختبر.", "منجز", "RLS والعزل والحذف المتسلسل"),
        ("ستة اختبارات ناجحة وبناء إنتاجي ناجح.", "منجز", "اختبارات العقد والبناء"),
        ("تم التحقق من عروض الهاتف الرئيسية دون تمرير أفقي.", "منجز", "التوافق المتجاوب"),
        ("تفعيلها من لوحة Auth.", "قبل النشر", "حماية كلمات المرور المسرّبة"),
        ("تدوير السر السابق ووضع السر الجديد في منصة الاستضافة فقط.", "قبل النشر", "تدوير مفاتيح Supabase"),
        ("تشغيل رحلة مصادقة كاملة من بيئة تملك اتصالًا خارجيًا.", "قبل استقبال بيانات", "E2E متصل"),
        ("تفعيل النسخ الاحتياطي واختبار الاستعادة.", "قبل استقبال بيانات", "استمرارية الأعمال"),
        ("اعتمادها من الجهة المشغلة وفق الولاية القضائية.", "حوكمة", "الخصوصية وشروط الاستخدام"),
    ], [4300, 1700, 3360], font_size=8.5)
    add_callout(doc, "التوصية النهائية", "الكود وقاعدة البيانات في حالة Release Candidate قابلة للنشر. فتح النظام لبيانات حساسة فعلية مشروط بإغلاق عناصر ما قبل النشر أعلاه، خصوصًا تدوير الأسرار، حماية كلمات المرور المسرّبة، النسخ الاحتياطي، واختبار رحلة المستخدم على رابط HTTPS النهائي.")
    doc.add_page_break()

    add_heading(doc, "ملحق أ - مصفوفة التتبع", 1)
    add_para(doc, "تربط المصفوفة مجموعات المتطلبات بمكونات التنفيذ واختبارات القبول. وهي مرجع التغيير عند تطوير إصدار لاحق.", color=MUTED)
    add_table(doc, ["اختبارات القبول", "مكونات التنفيذ", "المتطلبات"], [
        ("AT-004، AT-005، AT-009", "middleware، API auth، Supabase Auth", "FR-001 إلى FR-005"),
        ("AT-005، AT-006، AT-009", "SuperAdmin، admin routes، DB RPC", "FR-006 إلى FR-011"),
        ("AT-003، AT-009", "familyService، Families، Excel/Print", "FR-012 إلى FR-017"),
        ("AT-003، AT-009", "nominationService، Nominations، Excel/Print", "FR-018 إلى FR-023"),
        ("AT-003، AT-009", "Dashboard، Settings، announcements، renewal RPC", "FR-024 إلى FR-030"),
        ("AT-001 إلى AT-010", "Next config، RLS، build، monitoring/backup", "NFR-001 إلى NFR-010"),
    ], [2500, 4300, 2560], font_size=8.6)

    add_heading(doc, "ملحق ب - قائمة فحص التسليم", 1)
    add_list(doc, [
        "تثبيت متغيرات البيئة في منصة الاستضافة.",
        "تدوير المفتاح السري السابق وعدم نسخه إلى المتصفح أو المستودع.",
        "تفعيل حماية كلمات المرور المسرّبة وسياسة قوة كلمة المرور.",
        "نشر آخر نسخة بناء وربط النطاق الرسمي وHTTPS.",
        "اختبار المشرف العام ومدير المخيم على بيئة الإنتاج.",
        "اختبار إنشاء مخيم، إضافة أسرة وترشيح، تصدير كشف، ثم حذف المخيم التجريبي.",
        "تفعيل النسخ الاحتياطي واختبار الاستعادة وتوثيق المالك والزمن.",
        "اعتماد سياسة الخصوصية وشروط الاستخدام ووسيلة الدعم.",
    ])
    add_para(doc, "نهاية الوثيقة", size=9, color=MUTED, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=20)

    doc.core_properties.title = "مواصفات متطلبات البرمجيات - منصة نَسَق"
    doc.core_properties.subject = "SRS للإصدار 1.0 من نظام إدارة المخيمات"
    doc.core_properties.keywords = "SRS, Next.js, Supabase, إدارة المخيمات"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
